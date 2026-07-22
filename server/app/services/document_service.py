from io import BytesIO
from pathlib import Path
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Pt
from app.schemas.proposal_output import ProposalOutput

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
jinja_env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))


def generate_pdf(proposal: ProposalOutput, business_name: str) -> bytes:
    template = jinja_env.get_template("proposal.html")
    html_content = template.render(business_name=business_name, **proposal.model_dump())
    pdf_bytes = HTML(string=html_content).write_pdf()
    return pdf_bytes


def generate_docx(proposal: ProposalOutput, business_name: str) -> bytes:
    doc = Document()

    title = doc.add_heading("Business Proposal", level=0)
    subtitle = doc.add_paragraph(f"Prepared for {business_name}")
    subtitle.runs[0].italic = True

    sections = [
        ("Executive Summary", proposal.executive_summary),
        ("Scope of Work", proposal.scope_of_work),
        ("Timeline", proposal.timeline),
        ("Pricing Overview", proposal.pricing_overview),
        ("Next Steps", proposal.next_steps),
    ]

    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()